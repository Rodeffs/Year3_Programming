from docx import Document
from docx.shared import Cm


def main():
    document = Document("output.docx")

    document.add_picture("sample.png", width=Cm(5), height=Cm(5))
    document.add_paragraph("Подпись: логотип Arch Linux")
    
    document.save("output.docx")


if __name__ == "__main__":
    main()
